"""
Export functionality for meeting transcripts and summaries.
Supports exporting to Markdown, PDF, and Word formats.
"""

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple, Union

from fpdf import FPDF
import markdown
from docx import Document

from utils.logger import get_logger
from utils.config import config_manager

logger = get_logger("export")


class MeetingExporter:
    """Export meeting data to various formats."""
    
    def __init__(self):
        """Initialize the meeting exporter."""
        self.config = config_manager.config.storage
        self.default_export_dir = Path(self.config.export_directory)
        
        # Create the export directory if it doesn't exist
        self.default_export_dir.mkdir(parents=True, exist_ok=True)
    
    def export_to_markdown(self, meeting_data: Dict[str, Any], 
                          output_path: Optional[Path] = None) -> Path:
        """Export meeting data to a Markdown file.
        
        Args:
            meeting_data: Dictionary containing meeting data.
            output_path: Optional path to save the Markdown file.
            
        Returns:
            Path to the saved Markdown file.
        """
        try:
            # Create default output path if not provided
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                title_slug = self._slugify(meeting_data.get("title", "Meeting"))
                output_path = self.default_export_dir / f"{title_slug}_{timestamp}.md"
            
            # Extract data
            title = meeting_data.get("title", "Meeting Transcript")
            date = meeting_data.get("date", datetime.now())
            participants = meeting_data.get("participants", [])
            tags = meeting_data.get("tags", [])
            duration = meeting_data.get("duration", 0)
            summary = meeting_data.get("summary", {})
            transcript = meeting_data.get("transcript", {})
            action_items = meeting_data.get("action_items", [])
            
            # Format date
            date_str = date.strftime("%Y-%m-%d %H:%M") if isinstance(date, datetime) else str(date)
            
            # Format duration
            if duration:
                hours, remainder = divmod(int(duration), 3600)
                minutes, seconds = divmod(remainder, 60)
                duration_str = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
            else:
                duration_str = "Unknown"
            
            # Start building the Markdown content
            content = [
                f"# {title}",
                "",
                f"**Date:** {date_str}",
                f"**Duration:** {duration_str}",
                ""
            ]
            
            # Add participants if available
            if participants:
                content.append("## Participants")
                content.append("")
                for participant in participants:
                    content.append(f"- {participant}")
                content.append("")
            
            # Add tags if available
            if tags:
                content.append("**Tags:** " + ", ".join([f"`{tag}`" for tag in tags]))
                content.append("")
            
            # Add summary if available
            if summary:
                content.append("## Summary")
                content.append("")
                content.append(summary.get("summary_text", "No summary available."))
                content.append("")
                
                # Add key points if available
                key_points = summary.get("key_points", [])
                if key_points:
                    content.append("### Key Points")
                    content.append("")
                    for point in key_points:
                        content.append(f"- {point}")
                    content.append("")
                
                # Add topics if available
                topics = summary.get("topics", [])
                if topics:
                    content.append("### Topics Discussed")
                    content.append("")
                    for topic in topics:
                        topic_name = topic.get("name", "")
                        topic_discussion = topic.get("discussion", "")
                        content.append(f"#### {topic_name}")
                        content.append("")
                        content.append(topic_discussion)
                        content.append("")
                
                # Add decisions if available
                decisions = summary.get("decisions", [])
                if decisions:
                    content.append("### Decisions Made")
                    content.append("")
                    for decision in decisions:
                        content.append(f"- {decision}")
                    content.append("")
                
                # Add questions if available
                questions = summary.get("questions", [])
                if questions:
                    content.append("### Questions & Answers")
                    content.append("")
                    for qa in questions:
                        question = qa.get("question", "")
                        answer = qa.get("answer", "")
                        content.append(f"**Q:** {question}")
                        content.append(f"**A:** {answer}")
                        content.append("")
            
            # Add action items if available
            if action_items:
                content.append("## Action Items")
                content.append("")
                for item in action_items:
                    task = item.get("task", "")
                    assignee = item.get("assignee", "Unassigned")
                    due_date = item.get("due_date", "")
                    status = item.get("status", "pending")
                    
                    # Format due date
                    if isinstance(due_date, datetime):
                        due_date_str = due_date.strftime("%Y-%m-%d")
                    else:
                        due_date_str = str(due_date) if due_date else "Not specified"
                    
                    # Format status
                    status_icon = "✅" if status == "completed" else "⏳" if status == "pending" else "❌"
                    
                    content.append(f"- {status_icon} **{task}** - Assigned to: {assignee}, Due: {due_date_str}")
                
                content.append("")
            
            # Add transcript if available
            if transcript:
                content.append("## Full Transcript")
                content.append("")
                content.append("```")
                content.append(transcript.get("full_text", "No transcript available."))
                content.append("```")
            
            # Join all content and write to file
            markdown_content = "\n".join(content)
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(markdown_content)
            
            logger.info(f"Exported meeting to Markdown file: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting to Markdown: {e}")
            raise
    
    def export_to_pdf(self, meeting_data: Dict[str, Any], 
                     output_path: Optional[Path] = None) -> Path:
        """Export meeting data to a PDF file.
        
        Args:
            meeting_data: Dictionary containing meeting data.
            output_path: Optional path to save the PDF file.
            
        Returns:
            Path to the saved PDF file.
        """
        try:
            # Create default output path if not provided
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                title_slug = self._slugify(meeting_data.get("title", "Meeting"))
                output_path = self.default_export_dir / f"{title_slug}_{timestamp}.pdf"
            
            # Extract data
            title = meeting_data.get("title", "Meeting Transcript")
            date = meeting_data.get("date", datetime.now())
            participants = meeting_data.get("participants", [])
            tags = meeting_data.get("tags", [])
            duration = meeting_data.get("duration", 0)
            summary = meeting_data.get("summary", {})
            transcript = meeting_data.get("transcript", {})
            action_items = meeting_data.get("action_items", [])
            
            # Format date
            date_str = date.strftime("%Y-%m-%d %H:%M") if isinstance(date, datetime) else str(date)
            
            # Format duration
            if duration:
                hours, remainder = divmod(int(duration), 3600)
                minutes, seconds = divmod(remainder, 60)
                duration_str = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
            else:
                duration_str = "Unknown"
            
            # Create PDF
            pdf = FPDF()
            pdf.add_page()
            
            # Set font
            pdf.set_font("Arial", "B", 16)
            
            # Title
            pdf.cell(0, 10, title, 0, 1, "C")
            pdf.ln(10)
            
            # Metadata
            pdf.set_font("Arial", "", 12)
            pdf.cell(0, 10, f"Date: {date_str}", 0, 1)
            pdf.cell(0, 10, f"Duration: {duration_str}", 0, 1)
            
            # Participants
            if participants:
                pdf.ln(5)
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Participants", 0, 1)
                pdf.set_font("Arial", "", 12)
                
                for participant in participants:
                    pdf.cell(0, 10, f"- {participant}", 0, 1)
            
            # Tags
            if tags:
                pdf.ln(5)
                pdf.cell(0, 10, f"Tags: {', '.join(tags)}", 0, 1)
            
            # Summary
            if summary:
                pdf.ln(10)
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Summary", 0, 1)
                pdf.set_font("Arial", "", 12)
                
                # Split summary text into smaller chunks to avoid overflow
                summary_text = summary.get("summary_text", "No summary available.")
                for line in self._split_text_for_pdf(summary_text, 80):
                    pdf.multi_cell(0, 10, line)
                
                # Key points
                key_points = summary.get("key_points", [])
                if key_points:
                    pdf.ln(5)
                    pdf.set_font("Arial", "B", 13)
                    pdf.cell(0, 10, "Key Points", 0, 1)
                    pdf.set_font("Arial", "", 12)
                    
                    for point in key_points:
                        pdf.multi_cell(0, 10, f"- {point}")
                
                # Topics
                topics = summary.get("topics", [])
                if topics:
                    pdf.ln(5)
                    pdf.set_font("Arial", "B", 13)
                    pdf.cell(0, 10, "Topics Discussed", 0, 1)
                    
                    for topic in topics:
                        topic_name = topic.get("name", "")
                        topic_discussion = topic.get("discussion", "")
                        
                        pdf.set_font("Arial", "B", 12)
                        pdf.cell(0, 10, topic_name, 0, 1)
                        
                        pdf.set_font("Arial", "", 12)
                        for line in self._split_text_for_pdf(topic_discussion, 80):
                            pdf.multi_cell(0, 10, line)
                        
                        pdf.ln(5)
                
                # Decisions
                decisions = summary.get("decisions", [])
                if decisions:
                    pdf.ln(5)
                    pdf.set_font("Arial", "B", 13)
                    pdf.cell(0, 10, "Decisions Made", 0, 1)
                    pdf.set_font("Arial", "", 12)
                    
                    for decision in decisions:
                        pdf.multi_cell(0, 10, f"- {decision}")
            
            # Action items
            if action_items:
                pdf.add_page()
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Action Items", 0, 1)
                pdf.set_font("Arial", "", 12)
                
                for item in action_items:
                    task = item.get("task", "")
                    assignee = item.get("assignee", "Unassigned")
                    due_date = item.get("due_date", "")
                    status = item.get("status", "pending")
                    
                    # Format due date
                    if isinstance(due_date, datetime):
                        due_date_str = due_date.strftime("%Y-%m-%d")
                    else:
                        due_date_str = str(due_date) if due_date else "Not specified"
                    
                    # Format status
                    status_icon = "✓" if status == "completed" else "→" if status == "pending" else "✗"
                    
                    pdf.set_font("Arial", "B", 12)
                    pdf.multi_cell(0, 10, f"{status_icon} {task}")
                    
                    pdf.set_font("Arial", "", 12)
                    pdf.cell(0, 10, f"Assigned to: {assignee}", 0, 1)
                    pdf.cell(0, 10, f"Due: {due_date_str}", 0, 1)
                    pdf.ln(5)
            
            # Transcript (optional, can be very long)
            if transcript and transcript.get("full_text"):
                pdf.add_page()
                pdf.set_font("Arial", "B", 14)
                pdf.cell(0, 10, "Full Transcript", 0, 1)
                pdf.set_font("Arial", "", 10)  # Smaller font for transcript
                
                # Split transcript into smaller chunks to avoid overflow
                transcript_text = transcript.get("full_text", "No transcript available.")
                for line in self._split_text_for_pdf(transcript_text, 90):
                    pdf.multi_cell(0, 8, line)
            
            # Save PDF
            pdf.output(str(output_path))
            
            logger.info(f"Exported meeting to PDF file: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting to PDF: {e}")
            raise
    
    def export_to_docx(self, meeting_data: Dict[str, Any], 
                      output_path: Optional[Path] = None) -> Path:
        """Export meeting data to a Word document.
        
        Args:
            meeting_data: Dictionary containing meeting data.
            output_path: Optional path to save the Word document.
            
        Returns:
            Path to the saved Word document.
        """
        try:
            # Create default output path if not provided
            if not output_path:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                title_slug = self._slugify(meeting_data.get("title", "Meeting"))
                output_path = self.default_export_dir / f"{title_slug}_{timestamp}.docx"
            
            # Extract data
            title = meeting_data.get("title", "Meeting Transcript")
            date = meeting_data.get("date", datetime.now())
            participants = meeting_data.get("participants", [])
            tags = meeting_data.get("tags", [])
            duration = meeting_data.get("duration", 0)
            summary = meeting_data.get("summary", {})
            transcript = meeting_data.get("transcript", {})
            action_items = meeting_data.get("action_items", [])
            
            # Format date
            date_str = date.strftime("%Y-%m-%d %H:%M") if isinstance(date, datetime) else str(date)
            
            # Format duration
            if duration:
                hours, remainder = divmod(int(duration), 3600)
                minutes, seconds = divmod(remainder, 60)
                duration_str = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
            else:
                duration_str = "Unknown"
            
            # Create document
            doc = Document()
            
            # Add title
            doc.add_heading(title, level=1)
            
            # Add metadata
            doc.add_paragraph(f"Date: {date_str}")
            doc.add_paragraph(f"Duration: {duration_str}")
            
            # Add participants
            if participants:
                doc.add_heading("Participants", level=2)
                participants_para = doc.add_paragraph()
                for participant in participants:
                    participants_para.add_run(f"• {participant}\n")
            
            # Add tags
            if tags:
                tags_para = doc.add_paragraph("Tags: ")
                tags_para.add_run(", ".join(tags))
            
            # Add summary
            if summary:
                doc.add_heading("Summary", level=2)
                doc.add_paragraph(summary.get("summary_text", "No summary available."))
                
                # Add key points
                key_points = summary.get("key_points", [])
                if key_points:
                    doc.add_heading("Key Points", level=3)
                    for point in key_points:
                        doc.add_paragraph(point, style="List Bullet")
                
                # Add topics
                topics = summary.get("topics", [])
                if topics:
                    doc.add_heading("Topics Discussed", level=3)
                    for topic in topics:
                        topic_name = topic.get("name", "")
                        topic_discussion = topic.get("discussion", "")
                        
                        doc.add_heading(topic_name, level=4)
                        doc.add_paragraph(topic_discussion)
                
                # Add decisions
                decisions = summary.get("decisions", [])
                if decisions:
                    doc.add_heading("Decisions Made", level=3)
                    for decision in decisions:
                        doc.add_paragraph(decision, style="List Bullet")
                
                # Add questions
                questions = summary.get("questions", [])
                if questions:
                    doc.add_heading("Questions & Answers", level=3)
                    for qa in questions:
                        question = qa.get("question", "")
                        answer = qa.get("answer", "")
                        
                        q_para = doc.add_paragraph()
                        q_para.add_run("Q: ").bold = True
                        q_para.add_run(question)
                        
                        a_para = doc.add_paragraph()
                        a_para.add_run("A: ").bold = True
                        a_para.add_run(answer)
            
            # Add action items
            if action_items:
                doc.add_heading("Action Items", level=2)
                
                for item in action_items:
                    task = item.get("task", "")
                    assignee = item.get("assignee", "Unassigned")
                    due_date = item.get("due_date", "")
                    status = item.get("status", "pending")
                    
                    # Format due date
                    if isinstance(due_date, datetime):
                        due_date_str = due_date.strftime("%Y-%m-%d")
                    else:
                        due_date_str = str(due_date) if due_date else "Not specified"
                    
                    # Format status
                    status_map = {"completed": "✓ ", "pending": "→ ", "cancelled": "✗ "}
                    status_prefix = status_map.get(status, "")
                    
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(f"{status_prefix}{task}").bold = True
                    p.add_run(f"\nAssigned to: {assignee}")
                    p.add_run(f"\nDue: {due_date_str}")
            
            # Add transcript
            if transcript and transcript.get("full_text"):
                doc.add_page_break()
                doc.add_heading("Full Transcript", level=2)
                doc.add_paragraph(transcript.get("full_text", "No transcript available."))
            
            # Save document
            doc.save(str(output_path))
            
            logger.info(f"Exported meeting to Word document: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error exporting to Word: {e}")
            raise
    
    def _slugify(self, text: str) -> str:
        """Convert text to a URL-friendly slug.
        
        Args:
            text: Text to convert.
            
        Returns:
            Slugified text.
        """
        # Remove special characters
        slug = "".join(c if c.isalnum() or c in "-_" else "_" for c in text.lower())
        
        # Remove multiple consecutive underscores
        slug = "_".join(filter(None, slug.split("_")))
        
        return slug
    
    def _split_text_for_pdf(self, text: str, max_length: int = 80) -> List[str]:
        """Split text into chunks for PDF display.
        
        Args:
            text: Text to split.
            max_length: Maximum length of each chunk.
            
        Returns:
            List of text chunks.
        """
        # Replace newlines with spaces to avoid awkward line breaks
        text = text.replace("\n", " ").replace("\r", "")
        
        # Split into paragraphs
        paragraphs = text.split('\n\n')
        lines = []
        
        for paragraph in paragraphs:
            words = paragraph.split()
            current_line = []
            current_length = 0
            
            for word in words:
                word_length = len(word)
                
                if current_length + word_length + 1 <= max_length:
                    current_line.append(word)
                    current_length += word_length + 1
                else:
                    lines.append(" ".join(current_line))
                    current_line = [word]
                    current_length = word_length
            
            if current_line:
                lines.append(" ".join(current_line))
            
            # Add an empty line between paragraphs
            lines.append("")
        
        return lines